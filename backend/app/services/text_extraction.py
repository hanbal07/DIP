"""Text/document extraction for digital (non-scanned) files.

Handles digital PDFs (PyMuPDF), DOCX, TXT/MD, RTF and a fallback. For scanned PDFs where
text extraction yields little content, the pipeline routes to OCR instead.
"""
from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass

from app.core.config import settings

logger = logging.getLogger(__name__)


@dataclass
class ExtractedPage:
    number: int
    text: str


@dataclass
class ExtractionResult:
    pages: list[ExtractedPage]
    #: Mechanism actually used: "digital" | "ocr" | "mixed" | "none"
    method: str
    #: Mean extraction confidence (real signal for OCR; 1.0 for digital text).
    confidence: float = 0.0


def normalize_text(text: str) -> str:
    """Normalise whitespace and control characters for downstream processing."""
    if not text:
        return ""
    # Remove null bytes and other control chars except newline/tab.
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    # Collapse multiple blank lines.
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf_digital(pdf_bytes: bytes, max_pages: int | None = None) -> ExtractionResult:
    """Extract text from a digital PDF using PyMuPDF.

    Returns pages with flag indicating whether the page appears scanned (very little text).
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    max_pages = max_pages or settings.max_pages
    pages: list[ExtractedPage] = []
    total_pages = min(doc.page_count, max_pages)
    scanned_pages = 0
    for i in range(total_pages):
        page = doc.load_page(i)
        text = normalize_text(page.get_text("text"))
        pages.append(ExtractedPage(number=i + 1, text=text))
        if len(text.split()) < 5:
            scanned_pages += 1
    doc.close()

    total = len(pages)
    if total == 0:
        return ExtractionResult(pages=pages, method="none", confidence=0.0)
    scoratio = scanned_pages / total
    if scoratio >= 0.6:
        method = "ocr"
    elif scoratio > 0:
        method = "mixed"
    else:
        method = "digital"
    # Confidence: digital text is exact (1.0); OCR is filled in by the OCR step.
    conf = 1.0 if method == "digital" else 0.0
    return ExtractionResult(pages=pages, method=method, confidence=conf)


def extract_docx(docx_bytes: bytes) -> ExtractionResult:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(docx_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    # Include table text as best effort.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                paragraphs.append(" | ".join(cells))
    text = normalize_text("\n".join(paragraphs))
    return ExtractionResult(
        pages=[ExtractedPage(number=1, text=text)] if text else [],
        method="digital",
        confidence=1.0,
    )


def extract_plain(text: str) -> ExtractionResult:
    text = normalize_text(text)
    if not text:
        return ExtractionResult(pages=[], method="none", confidence=0.0)
    return ExtractionResult(
        pages=[ExtractedPage(number=1, text=text)], method="digital", confidence=1.0
    )
